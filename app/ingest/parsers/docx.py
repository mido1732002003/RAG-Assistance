"""Word document parser."""

from pathlib import Path
from typing import Dict, Any
import docx

from app.ingest.parsers.base import BaseParser
from app.utils.logging import get_logger

logger = get_logger(__name__)


class DocxParser(BaseParser):
    """Parse Word documents."""
    
    async def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX file."""
        if not self.validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        
        try:
            doc = docx.Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine all text
            text = "\n\n".join(paragraphs)
            if tables_text:
                text += "\n\n" + "\n\n".join(tables_text)
            
            # Extract metadata
            metadata = {}
            if doc.core_properties:
                metadata = {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                }
            
            return {
                "text": text,
                "metadata": metadata,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "title": metadata.get("title") or file_path.stem,
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise